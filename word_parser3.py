from docx import Document
from docx.shared import Inches
import docx2txt
import os
import sys
os.mkdir('unformatted_text_files')
os.mkdir('word_parser_docx_file')
os.mkdir('formatted_text_files')

def main():
    document_read = raw_input("Enter file name or path\n")
    document = Document(document_read)
    num = 1
    for para in document.paragraphs:
        name = "Topic" + str(num -1) + ".docx"
        name1 = "Topic" + str(num - 1) + ".txt"
        name3 = os.path.join('word_parser_docx_file', name)
        name4 = os.path.join('unformatted_text_files', name1)
        if para.style.name == 'Heading 1' and num == 1:
            doc = Document()
            sty = int(para.style.name[8])
            doc.add_heading("/*" + para.text + "*/", 1)
            num = num + 1
        elif para.style.name == 'Heading 1' and num > 1:
            doc.save(os.path.join('word_parser_docx_file', name))
            text = docx2txt.process(name3)
            f = open(name4, 'w')
            f.write(text)
            f.close()
            rc(name4, name1)
            doc = Document()
            doc.add_heading("/*" + para.text + "*/", sty)
            num = num + 1
        elif para.style.name != 'Heading 1' and para.style.name.startswith('Heading'):
            sty = int(para.style.name[8])
            doc.add_heading("/*" + para.text + "*/", sty)
        else:
            doc.add_paragraph(para.text)
    doc.save(os.path.join('word_parser_docx_file', name))
    text1 = docx2txt.process(name3)
    f1 = open(name4, 'w')
    f1.write(text1)
    f1.close()
    rc(name4, name1)

def rc(name1, name2):
    f = open(name1, "r")
    f1 = open(os.path.join('formatted_text_files', name2),"w")
    for lines in f:
        if lines != '\n':
            f1.write(lines)
    f.close()
    f1.close()
main()
